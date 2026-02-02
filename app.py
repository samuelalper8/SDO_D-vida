import streamlit as st
import fitz  # PyMuPDF
import pandas as pd
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import zipfile
from datetime import datetime

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Gerador DFF - ConPrev", layout="wide", page_icon="📄")

# --- FUNÇÕES DE EXTRAÇÃO (Mantida a lógica que já funcionava) ---
def extrair_dados_pdf(pdf_stream, filename):
    doc = fitz.open(stream=pdf_stream, filetype="pdf")
    texto_completo = ""
    for pagina in doc:
        texto_completo += pagina.get_text()
    
    # Regex ajustados
    municipio_match = re.search(r"MUNICIPIO DE\s+(.*)", texto_completo)
    municipio = municipio_match.group(1).strip() if municipio_match else "DESCONHECIDO"
    
    cnpj_match = re.search(r"(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2})", texto_completo)
    cnpj = cnpj_match.group(1) if cnpj_match else ""

    parcelamentos = []
    linhas = [l.strip() for l in texto_completo.split('\n') if l.strip()]
    
    for i, linha in enumerate(linhas):
        if cnpj and linha.startswith(cnpj):
            match_processo = re.search(r"\s+(\d{9,})", linha)
            match_valor = re.search(r"(\d{1,3}(?:\.\d{3})*,\d{2})", linha)
            
            if match_processo:
                valor = match_valor.group(1) if match_valor else "Verificar"
                if not match_valor and i + 1 < len(linhas):
                    prox_valor = re.search(r"(\d{1,3}(?:\.\d{3})*,\d{2})", linhas[i+1])
                    if prox_valor: valor = prox_valor.group(1)
                
                parcelamentos.append({"Processo": match_processo.group(1), "Saldo": valor})

    if not parcelamentos:
        saldo_total = re.search(r"SALDO DEVEDOR TOTAL\s+([\d.,]+)", texto_completo)
        valor_total = saldo_total.group(1) if saldo_total else "0,00"
        parcelamentos.append({"Processo": "Consolidado", "Saldo": valor_total})

    return {"Arquivo": filename, "Município": municipio, "CNPJ": cnpj, "Parcelamentos": parcelamentos}

# --- FUNÇÃO DE GERAÇÃO DO WORD (REFINADA) ---
def gerar_docx(dados, nome_prefeito, img_timbrado, img_sig_rubens, img_sig_samuel):
    doc = Document()
    
    # 1. Configurar Margens para acomodar o Papel Timbrado
    section = doc.sections[0]
    section.top_margin = Cm(3)    # Margem maior para o logo
    section.bottom_margin = Cm(3) # Margem maior para o rodapé
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 2. Inserir Papel Timbrado no Cabeçalho (Como imagem de fundo)
    if img_timbrado:
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(img_timbrado, width=Cm(21.5)) # Largura total da A4
        # Ajustes para tentar centralizar/posicionar (Python-docx tem limites aqui, mas o header ajuda)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.space_after = Pt(0)

    # 3. Estilo de Fonte Global (Arial 11/12)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # 4. Data
    meses = {1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho", 
             7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"}
    dt = datetime.now()
    data_extenso = f"Goiânia, {dt.day} de {meses[dt.month]} de {dt.year}."
    
    p_data = doc.add_paragraph(data_extenso)
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph() # Espaço

    # 5. Destinatário (Negrito e Caixa Alta conforme imagem)
    p_dest = doc.add_paragraph()
    runner = p_dest.add_run("EXCELENTÍSSIMO SENHOR\n")
    runner.bold = True
    runner = p_dest.add_run(f"{nome_prefeito.upper()}\n")
    runner.bold = True
    p_dest.add_run(f"PREFEITO MUNICIPAL DE {dados['Município'].upper()} – GO")

    doc.add_paragraph() 

    # 6. Assunto e Corpo (Texto Exato do Modelo)
    p_assunto = doc.add_paragraph()
    p_assunto.add_run("Assunto: ").bold = True
    p_assunto.add_run("Ficam apresentados os valores e a documentação comprobatória dos saldos de débitos existentes em 31 de dezembro de 2025, destinados à composição do Balanço Patrimonial.")
    p_assunto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("Senhor Prefeito,")

    texto_corpo = (
        "Ao tempo em que lhe cumprimento, na qualidade de assessoria do Município para assuntos relacionados a atos de pessoal e ao fisco federal, "
        "no âmbito das ações de conformidade administrativa, venho, por meio do presente, apresentar os valores e a documentação requisitados por esta assessoria especializada, "
        "referentes aos saldos de débitos destinados à composição do Balanço Patrimonial."
    )
    p_corpo = doc.add_paragraph(texto_corpo)
    p_corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_corpo.paragraph_format.first_line_indent = Cm(1.25) # Parágrafo tradicional

    doc.add_paragraph("Nesse contexto, discriminam-se abaixo o órgão de origem, o número do processo e os respectivos valores apurados em 31/12/2025, estando anexa a documentação comprobatória pertinente:")

    # 7. Tabela Formatada
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid' # Adiciona as bordas simples
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Órgão'
    hdr_cells[1].text = 'Processo / Documento'
    hdr_cells[2].text = 'Saldo em 31/12/2025'
    
    # Negrito no header da tabela
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)

    # Dados RFB
    for p in dados['Parcelamentos']:
        row = table.add_row().cells
        row[0].text = "Receita Federal do Brasil"
        row[1].text = str(p['Processo'])
        row[2].text = str(p['Saldo'])

    # Dados PGFN (Vazio)
    row_pgfn = table.add_row().cells
    row_pgfn[0].text = "Procuradoria da Fazenda Nacional"
    row_pgfn[1].text = "-"
    row_pgfn[2].text = "-"

    doc.add_paragraph()

    # 8. Texto Final
    texto_final = (
        "Solicita-se, por oportuno, que a referida documentação seja encaminhada ao setor contábil, a fim de que sejam adotadas as providências e registros contábeis cabíveis, conforme as normas aplicáveis.\n"
        "Esta consultoria agradece a confiança depositada e permanece à disposição para quaisquer esclarecimentos adicionais que se fizerem necessários.\n"
        "Sem mais para o momento, reitero protestos de elevada estima e consideração."
    )
    p_final = doc.add_paragraph(texto_final)
    p_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_final.paragraph_format.first_line_indent = Cm(1.25)
    
    doc.add_paragraph("Atenciosamente,")
    doc.add_paragraph() # Espaço extra antes das assinaturas

    # 9. Assinaturas (Layout com Tabela Invisível para alinhar imagens e texto)
    # Criamos uma tabela 1x3 (Rubens | Glayzer | Samuel) ou conforme a imagem (Vertical ou Horizontal)
    # A imagem mostra uma disposição vertical centralizada. Vamos seguir isso.
    
    # --- Rubens ---
    p_rubens = doc.add_paragraph()
    p_rubens.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img_sig_rubens:
        run_sig = p_rubens.add_run()
        run_sig.add_picture(img_sig_rubens, width=Cm(4)) # Tamanho da assinatura
        run_sig.add_break()
    
    run_nome = p_rubens.add_run("Rubens Pires Malaquias\n")
    run_nome.bold = True
    p_rubens.add_run("Diretor Técnico e Consultor junto ao Fisco Federal\nCRA/GO 6-007-48")
    
    doc.add_paragraph() # Espaço entre assinaturas

    # --- Glayzer ---
    p_glayzer = doc.add_paragraph()
    p_glayzer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Sem imagem de assinatura para Glayzer no modelo, apenas texto
    run_nome = p_glayzer.add_run("Glayzer Antônio Gomes da Silva\n")
    run_nome.bold = True
    p_glayzer.add_run("Advogado Especialista em Direito Público\nOAB/GO 28.315")

    doc.add_paragraph()

    # --- Samuel ---
    p_samuel = doc.add_paragraph()
    p_samuel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img_sig_samuel:
        run_sig = p_samuel.add_run()
        run_sig.add_picture(img_sig_samuel, width=Cm(4))
        run_sig.add_break()

    run_nome = p_samuel.add_run("Samuel Almeida\n")
    run_nome.bold = True
    p_samuel.add_run("Líder de Equipe – Fisco Federal GO")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFACE LATERAL (Sidebar) ---
st.sidebar.header("🛠️ Configurações Visuais")
st.sidebar.markdown("Para o documento sair perfeito, faça upload dos arquivos abaixo:")
f_timbrado = st.sidebar.file_uploader("1. Papel Timbrado (JPG/PNG)", type=["jpg", "png", "jpeg"])
f_sig_rubens = st.sidebar.file_uploader("2. Assinatura Rubens (PNG Transparente)", type=["png", "jpg"])
f_sig_samuel = st.sidebar.file_uploader("3. Assinatura Samuel (PNG Transparente)", type=["png", "jpg"])

# --- INTERFACE PRINCIPAL ---
st.title("📄 Gerador de Ofícios - Padrão ConPrev")
st.markdown("**Passo 1:** Arraste os PDFs de 'Saldo da Dívida RFB' abaixo.")

uploaded_files = st.file_uploader("Arquivos PDF", type="pdf", accept_multiple_files=True)

if uploaded_files:
    # Processamento
    lista_dados = []
    with st.spinner("Lendo PDFs..."):
        for f in uploaded_files:
            dados = extrair_dados_pdf(f.read(), f.name)
            lista_dados.append(dados)
    
    # Edição de Prefeitos
    st.markdown("**Passo 2:** Informe o nome dos Prefeitos.")
    df_preparacao = []
    for d in lista_dados:
        df_preparacao.append({
            "Arquivo": d["Arquivo"],
            "Município": d["Município"],
            "Prefeito (Preencher)": "A/C GESTOR MUNICIPAL" # Padrão
        })
    
    df_editado = st.data_editor(
        pd.DataFrame(df_preparacao),
        column_config={
            "Prefeito (Preencher)": st.column_config.TextColumn("Nome do Prefeito (Opcional)", help="Se vazio, usará 'A/C Gestor Municipal'")
        },
        use_container_width=True,
        hide_index=True
    )

    if st.button("Gerar Ofícios (.docx) com Timbre e Assinaturas"):
        if not f_timbrado:
            st.warning("⚠️ Você não anexou o Papel Timbrado na barra lateral. O fundo sairá branco.")
        
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            progress = st.progress(0)
            for i, row in df_editado.iterrows():
                dados_originais = next(d for d in lista_dados if d["Arquivo"] == row["Arquivo"])
                
                # Prepara imagens (reseta ponteiro de leitura para cada loop)
                img_t = BytesIO(f_timbrado.getvalue()) if f_timbrado else None
                img_r = BytesIO(f_sig_rubens.getvalue()) if f_sig_rubens else None
                img_s = BytesIO(f_sig_samuel.getvalue()) if f_sig_samuel else None
                
                docx = gerar_docx(
                    dados_originais, 
                    row["Prefeito (Preencher)"],
                    img_t, img_r, img_s
                )
                
                nome_arq = f"OFICIO_{row['Município'].replace(' ', '_').upper()}.docx"
                zf.writestr(nome_arq, docx.getvalue())
                progress.progress((i + 1) / len(df_editado))
        
        zip_buffer.seek(0)
        st.success("Documentos gerados com sucesso!")
        st.download_button("⬇️ Baixar Lote Completo (ZIP)", zip_buffer, "Oficios_ConPrev_2025.zip", "application/zip")
